from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

doc = Document()

# Set A4 LANDSCAPE orientation with small margins
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)  # A4 landscape width
section.page_height = Cm(21.0)  # A4 landscape height
section.left_margin = Cm(0.5)
section.right_margin = Cm(0.5)
section.top_margin = Cm(0.5)
section.bottom_margin = Cm(0.5)

# SET BEIGE BACKGROUND COLOR FOR THE PAGE - #f5f0e6
bg = parse_xml(f'<w:background {nsdecls("w")} w:color="F5F0E6"/>')
doc.element.insert(0, bg)
settings = doc.settings.element
display_bg = parse_xml(f'<w:displayBackgroundShape {nsdecls("w")}/>')
settings.append(display_bg)

# Title - using paragraph NOT heading (no blue underline)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('Diário de Sono')
title_run.font.size = Pt(18)
title_run.font.name = 'Arial'
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph()  # Small space

# Helper function to set cell shading
def set_cell_shading(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# Helper to set font for cell
def format_cell(cell, text, bold=False, size=8, italic=False, bg_color=None):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = 'Arial'
            run.bold = bold
            run.italic = italic
            run.font.color.rgb = RGBColor(0, 0, 0)
    if bg_color:
        set_cell_shading(cell, bg_color)

# Create single table - 20 rows x 8 columns
table = doc.add_table(rows=20, cols=8)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths to match PDF proportions
col_widths = [Cm(6.0), Cm(3.0), Cm(3.5), Cm(3.0), Cm(3.2), Cm(3.2), Cm(3.0), Cm(3.0)]
for row in table.rows:
    for i, width in enumerate(col_widths):
        row.cells[i].width = width

# Column headers - centered
col_headers = ['', '4ªf\n7/1/26', '5ªf\n8/1/26', '6ªf\n9/1/26', 'Sábado\n10/1/26', 'Domingo\n11/1/26', '2ªf\n12/1/26', '3ªf\n13/1/26']

# Row 0: Headers
for i, header in enumerate(col_headers):
    cell = table.cell(0, i)
    cell.text = header
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.name = 'Arial'
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    set_cell_shading(cell, 'F0F0F0')

# All data rows - matching PDF exactly
all_rows = [
    # MANHÃ section header (row 1)
    ('section', 'MANHÃ'),
    # MANHÃ data (rows 2-4)
    ('data', ['A que horas acordou esta manhã?', '15.00', '8h30 (fui c/ carmo ao hospital)', '15h.', '12:30m', '06.45m.', '15h.', '14:30']),
    ('data', ['A que horas se levantou da cama esta manhã?', '15.30', '8.40m', '15:30 Fui farmácia', '13h15m', '11h.00 (brunch da carmo)', '15.30', '15:00 mãe insistiu']),
    ('data', ['Como se sente esta manhã?\n(Muito mal / Mal / Razoavelmente / Bem / Muito bem)', 'Mal ✗', 'Mal ✗', 'Mal ✗', 'Mal ✗', 'Mal ✗', 'Mal ✗', 'Razoavelmente ✗']),
    # FATORES DIÁRIOS section header (row 5)
    ('section', 'FATORES DIÁRIOS'),
    # FATORES DIÁRIOS data (rows 6-9)
    ('data', ['Sestas durante o dia (Sim/Não)\n- Se sim, duração:', '—', '—', '—', 'Muita', '14/15h: (carro) 30min\n16:30h (rede museu) 15min', '—', '—']),
    ('data', ['Ingestão de Cafeína (chás, cafés, etc)\n- Tipo e quantidade', '—', 'Sim', '1.', '1 - amoreiras', '1 - (amoreiras)', '—', '—']),
    ('data', ['- Última hora de Consumo', '—', '2 Bica - 14h', '1 Bica 14h', '—', '15.30m', '—', '—']),
    ('data', ['Atividade Física (tipo e duração)', '—', '—', '—', 'Caminhada', 'Bicicleta 10m\n(Amoreiras à Gulbenkian)', 'Única refeição foi o jantar', '—']),
    # NOITE section header (row 10)
    ('section', 'NOITE'),
    # NOITE data (rows 11-19)
    ('data', ['A que horas se deitou na noite passada?', '4h00\n1.00', '5h. depois de ter bebido 1 copo c amigo', '5h', '5h.00\n(retiro a v. ritmo)', '5.30', '1h am', '—']),
    ('data', ['Quanto tempo demorou a adormecer? (em minutos)', 'logo pq estava cheio de sono', 'logo pq estava cheio de sono', 'logo pq estava cheio de sono', 'logo pq estava cheio de sono', 'logo pq estava cheio de sono', 'logo pq estava cheio de sono', '—']),
    ('data', ['Quantas vezes acordou durante a noite?', '—', '—', '—', '—', '—', '—', '—']),
    ('data', ['O que fez durante os despertares noturnos?', '—', '—', '—', '—', '—', '—', '—']),
    ('merged_italic', ['Quanto tempo esteve acordado durante a noite?\n(total em minutos)', 'Não acordei a meio da noite, mas estive acordado até à hora que escrevi deitar-me na célula acima']),
    ('data', ['Quanto tempo dormiu ao todo? (em horas/minutos)', '4h\n(04:00-08:30)', '10h.\n(05:00-15:00)', '8h30\n(04:30-12:15)', '6h45\n(05:00-10:45)', '9h30\n(05:30-15:00)', '11h\n(03:00-14:30)', '—']),
    ('data', ['Que comprimidos tomou para dormir na noite\npassada? Quantos?', '—', '—', '—', '—', '—', '—', '—']),
    ('data', ['Que quantidade de bebidas alcoólicas ingeriu na\nnoite passada?', '—', '2 cervejas', 'Cappriciosa jantar, mas não bebi', 'Festa + jantar lx\n3 cervejas', '—', '—', '—']),
    ('data', ['Como acha que passou a noite?\n(Muito mal / Mal / Razoavelmente / Bem / Muito bem)', 'Muito bem ✗', 'Muito bem ✗', 'Muito bem ✗', 'Muito bem ✗', 'Muito bem ✗', 'Muito bem ✗', 'Muito bem ✗']),
]

row_idx = 1
for row_type, row_content in all_rows:
    if row_type == 'section':
        # Merge all cells for section header
        cell = table.cell(row_idx, 0)
        cell.merge(table.cell(row_idx, 7))
        format_cell(cell, row_content, bold=True, size=9, bg_color='E8E8E8')
    elif row_type == 'merged_italic':
        # Row label
        format_cell(table.cell(row_idx, 0), row_content[0], bold=True, size=8, bg_color='FFFFFF')
        # Merge cells 1-7 for italic text
        merged_cell = table.cell(row_idx, 1)
        merged_cell.merge(table.cell(row_idx, 7))
        format_cell(merged_cell, row_content[1], italic=True, size=8, bg_color='FFFFFF')
    else:
        for col_idx, cell_text in enumerate(row_content):
            cell = table.cell(row_idx, col_idx)
            if col_idx == 0:
                format_cell(cell, cell_text, bold=True, size=8, bg_color='FFFFFF')
            else:
                format_cell(cell, cell_text, bold=False, size=8, bg_color='FFFFFF')
    row_idx += 1

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer.add_run('Gerado a 13 de Janeiro de 2026')
footer_run.font.size = Pt(8)
footer_run.font.name = 'Arial'
footer_run.font.color.rgb = RGBColor(0, 0, 0)

doc.save('/Users/tomasbatalha/Downloads/DIARIO_SONO_7-13_JAN.docx')
print("Word document created - NO BLUE LINE, MATCHING PDF!")
